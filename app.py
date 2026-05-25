import hashlib
import io
import os
import re
import textwrap
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import fitz
import pandas as pd
import pdfplumber
import plotly.express as px
import plotly.graph_objects as go
import requests
import streamlit as st
from docx import Document
from tinydb import Query, TinyDB

try:
    from bs4 import BeautifulSoup
except Exception:
    BeautifulSoup = None

try:
    from groq import Groq
except Exception:
    Groq = None

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from langsmith import traceable
except Exception:
    traceable = None

try:
    import textstat
except Exception:
    textstat = None

try:
    from langdetect import detect
except Exception:
    detect = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except Exception:
    TfidfVectorizer = None
    cosine_similarity = None


APP_DIR = Path(__file__).parent
DATA_DIR = APP_DIR / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
DB_PATH = DATA_DIR / "history.json"
DATA_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

st.set_page_config(
    page_title="AI Content Summarizer",
    page_icon="AI",
    layout="wide",
    initial_sidebar_state="expanded",
)


CSS = """
<style>
:root {
  --bg: #080b10;
  --panel: rgba(18, 24, 35, .84);
  --panel-2: rgba(24, 32, 46, .72);
  --line: rgba(255,255,255,.10);
  --text: #eef4ff;
  --muted: #95a2b8;
  --aqua: #36d1dc;
  --lime: #a7f070;
  --rose: #ff719a;
  --gold: #f4c95d;
}
.stApp {
  background:
    radial-gradient(circle at 15% 8%, rgba(54, 209, 220, .20), transparent 32%),
    radial-gradient(circle at 85% 12%, rgba(167, 240, 112, .13), transparent 28%),
    linear-gradient(135deg, #07090e 0%, #101521 48%, #090d14 100%);
  color: var(--text);
}
section[data-testid="stSidebar"] {
  background: rgba(7, 10, 16, .90);
  border-right: 1px solid var(--line);
}
div[data-testid="stSidebarUserContent"] {
  padding-top: 1.25rem;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div[data-testid="stMarkdownContainer"] {
  color: #e7eefc !important;
}
section[data-testid="stSidebar"] small,
section[data-testid="stSidebar"] [data-testid="stCaptionContainer"],
section[data-testid="stSidebar"] .stRadio p {
  color: #aebad0 !important;
}
section[data-testid="stSidebar"] [role="radiogroup"] label p {
  color: #c8d3e6 !important;
}
section[data-testid="stSidebar"] [role="radiogroup"] label:has(input:checked) p {
  color: #ffffff !important;
  font-weight: 700;
}
label,
.stMarkdown,
.stTextInput label,
.stTextArea label,
.stSelectbox label,
.stFileUploader label,
.stRadio label {
  color: #dce7f8 !important;
}
.stRadio [role="radiogroup"] label p,
.stRadio [role="radiogroup"] label span,
div[role="radiogroup"] label p,
div[role="radiogroup"] label span {
  color: #dbe6f7 !important;
  opacity: 1 !important;
}
.stRadio [role="radiogroup"] label:has(input:checked) p,
div[role="radiogroup"] label:has(input:checked) p {
  color: #ffffff !important;
  font-weight: 700;
}
.stTextInput input,
.stTextArea textarea,
.stSelectbox [data-baseweb="select"] {
  color: #0d1522 !important;
}
.hero {
  padding: 1.1rem 1.35rem;
  border: 1px solid var(--line);
  border-radius: 8px;
  background: linear-gradient(120deg, rgba(54,209,220,.13), rgba(167,240,112,.07)), rgba(13,18,28,.78);
  box-shadow: 0 18px 60px rgba(0,0,0,.28);
  animation: rise .45s ease-out both;
}
.hero h1 {
  margin: 0 0 .25rem 0;
  font-size: clamp(2rem, 4vw, 4rem);
  letter-spacing: 0;
  line-height: 1;
}
.hero p {
  color: var(--muted);
  margin: .25rem 0 0 0;
  max-width: 980px;
}
.metric-card, .glass {
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--panel);
  box-shadow: 0 16px 48px rgba(0,0,0,.22);
  padding: 1rem;
}
.metric-card {
  min-height: 116px;
  animation: fadeIn .5s ease-out both;
  overflow-wrap: normal;
  word-break: normal;
}
.metric-label { color: var(--muted); font-size: .83rem; }
.metric-value { font-size: 1.55rem; font-weight: 760; margin-top: .2rem; line-height: 1.2; }
.metric-note { color: var(--muted); font-size: .8rem; margin-top: .15rem; }
.badge {
  display: inline-flex;
  align-items: center;
  gap: .35rem;
  border: 1px solid var(--line);
  border-radius: 999px;
  padding: .28rem .62rem;
  color: #dce8ff;
  background: rgba(255,255,255,.055);
  font-size: .82rem;
  margin: .1rem .25rem .1rem 0;
}
.summary-box {
  background: rgba(255,255,255,.045);
  border-left: 3px solid var(--aqua);
  border-radius: 8px;
  padding: 1rem 1.1rem;
  line-height: 1.65;
}
.drop-hint {
  border: 1px dashed rgba(255,255,255,.22);
  background: rgba(255,255,255,.035);
  border-radius: 8px;
  padding: .85rem;
  color: #bdc9dc;
  line-height: 1.55;
  margin: .45rem 0 .85rem 0;
}
.stButton>button, .stDownloadButton>button {
  border-radius: 8px;
  border: 1px solid rgba(255,255,255,.14);
  background: linear-gradient(135deg, #2fc9d6, #85db78);
  color: #071014;
  font-weight: 700;
  min-height: 3rem;
  white-space: normal;
}
.stButton>button p, .stDownloadButton>button p,
.stButton>button span, .stDownloadButton>button span {
  color: #071014 !important;
  font-weight: 700;
}
.stTabs [data-baseweb="tab-list"] { gap: .45rem; }
.stTabs [data-baseweb="tab"] {
  border-radius: 8px;
  background: rgba(255,255,255,.055);
  border: 1px solid var(--line);
  padding: .45rem .8rem;
}
.stTabs [data-baseweb="tab"] p,
.stTabs [data-baseweb="tab"] span {
  color: #dce7f8 !important;
  opacity: 1 !important;
}
.stTabs [aria-selected="true"] {
  background: rgba(54,209,220,.18);
  border-color: rgba(54,209,220,.45);
}
.stTabs [aria-selected="true"] p,
.stTabs [aria-selected="true"] span {
  color: #ffffff !important;
  font-weight: 700;
}
div[data-testid="stFileUploader"] section {
  border-radius: 8px;
  background: rgba(255,255,255,.04);
  border: 1px dashed rgba(54,209,220,.35);
}
div[data-testid="stFileUploader"] section * {
  color: #dce7f8 !important;
}
div[data-testid="stFileUploader"] section button {
  background: linear-gradient(135deg, #2fc9d6, #85db78) !important;
  border: 1px solid rgba(255,255,255,.18) !important;
  color: #071014 !important;
  opacity: 1 !important;
  box-shadow: 0 10px 28px rgba(54, 209, 220, .18);
}
div[data-testid="stFileUploader"] section button * {
  color: #071014 !important;
  fill: #071014 !important;
}
div[data-testid="stFileUploader"] section button:hover {
  filter: brightness(1.06);
  transform: translateY(-1px);
}
div[data-testid="stFileUploader"] small {
  color: #9fb0c8 !important;
}
@keyframes rise { from { opacity: 0; transform: translateY(12px); } to { opacity: 1; transform: translateY(0); } }
@keyframes fadeIn { from { opacity: 0; transform: translateY(8px); } to { opacity: 1; transform: translateY(0); } }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)


STOPWORDS = {
    "the", "and", "for", "that", "with", "this", "from", "are", "was", "were", "have", "has",
    "had", "not", "but", "you", "your", "our", "their", "they", "them", "his", "her", "its",
    "into", "than", "then", "there", "these", "those", "about", "would", "could", "should",
    "will", "can", "may", "also", "more", "most", "such", "when", "where", "which", "who",
    "what", "how", "why", "been", "being", "over", "under", "between", "within", "while",
}

POSITIVE = {"good", "great", "excellent", "strong", "positive", "growth", "improve", "improved", "success", "benefit", "clear", "effective", "efficient", "opportunity", "innovative"}
NEGATIVE = {"bad", "weak", "negative", "risk", "risks", "decline", "problem", "issue", "issues", "fail", "failed", "failure", "threat", "poor", "costly", "unclear", "difficult"}


def db() -> TinyDB:
    return TinyDB(DB_PATH)


def now_label() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def clean_text(text: str) -> str:
    text = re.sub(r"\r", "\n", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def sentences(text: str) -> List[str]:
    parts = re.split(r"(?<=[.!?])\s+|\n+", clean_text(text))
    return [p.strip() for p in parts if len(p.strip()) > 25]


def words(text: str) -> List[str]:
    return re.findall(r"[A-Za-z][A-Za-z'-]{2,}", text.lower())


def chunks(text: str, max_chars: int = 9000) -> List[str]:
    source = clean_text(text)
    if len(source) <= max_chars:
        return [source]
    result, buff = [], []
    size = 0
    for sent in sentences(source):
        if size + len(sent) > max_chars and buff:
            result.append(" ".join(buff))
            buff, size = [], 0
        buff.append(sent)
        size += len(sent) + 1
    if buff:
        result.append(" ".join(buff))
    return result[:8]


def extract_keywords(text: str, limit: int = 18) -> List[Tuple[str, int]]:
    tokens = [w for w in words(text) if w not in STOPWORDS and len(w) > 3]
    return Counter(tokens).most_common(limit)


def extractive_summary(text: str, max_sentences: int = 5) -> str:
    sents = sentences(text)
    if not sents:
        return clean_text(text)[:900]
    keywords = dict(extract_keywords(text, 80))
    scored = []
    for idx, sent in enumerate(sents):
        score = sum(keywords.get(w, 0) for w in words(sent))
        score += max(0, 6 - idx) * 1.5
        scored.append((score, idx, sent))
    selected = sorted(sorted(scored, reverse=True)[:max_sentences], key=lambda x: x[1])
    return " ".join(item[2] for item in selected)


def detect_language(text: str) -> str:
    if not text:
        return "Unknown"
    if detect:
        try:
            return detect(text[:2000]).upper()
        except Exception:
            return "Unknown"
    return "Unknown"


def readability(text: str) -> Dict[str, float]:
    if textstat:
        try:
            return {
                "Flesch Reading Ease": round(textstat.flesch_reading_ease(text), 1),
                "Grade Level": round(textstat.flesch_kincaid_grade(text), 1),
                "Reading Time": round(len(words(text)) / 220, 1),
            }
        except Exception:
            pass
    word_count = max(1, len(words(text)))
    sent_count = max(1, len(sentences(text)))
    avg_sentence = word_count / sent_count
    return {
        "Flesch Reading Ease": round(max(0, 100 - avg_sentence * 1.7), 1),
        "Grade Level": round(max(1, avg_sentence / 2.2), 1),
        "Reading Time": round(word_count / 220, 1),
    }


def sentiment(text: str) -> Dict[str, float]:
    token_list = words(text)
    total = max(1, len(token_list))
    pos = sum(1 for w in token_list if w in POSITIVE)
    neg = sum(1 for w in token_list if w in NEGATIVE)
    score = (pos - neg) / max(1, pos + neg)
    label = "Positive" if score > 0.15 else "Negative" if score < -0.15 else "Neutral"
    return {"label": label, "score": round(score, 2), "positive": pos, "negative": neg, "density": round((pos + neg) / total, 3)}


def topic_breakdown(text: str, topic_count: int = 5) -> List[Dict[str, str]]:
    sents = sentences(text)
    if not sents:
        return []
    kws = [kw for kw, _ in extract_keywords(text, 60)]
    topics = []
    used = set()
    for keyword in kws:
        matches = [s for s in sents if keyword in s.lower() and s not in used]
        if matches:
            used.update(matches[:2])
            topics.append({"topic": keyword.title(), "summary": extractive_summary(" ".join(matches[:5]), 2)})
        if len(topics) >= topic_count:
            break
    return topics


def extract_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
    except Exception:
        pass
    if not clean_text("\n".join(text_parts)):
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
        except Exception:
            pass
    return clean_text("\n".join(text_parts))


def extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return clean_text("\n".join(parts))


def extract_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return clean_text(file_bytes.decode(enc))
        except Exception:
            continue
    return ""


def fetch_webpage(url: str) -> Tuple[str, str]:
    if not url:
        return "", ""
    if not BeautifulSoup:
        raise RuntimeError("BeautifulSoup is not installed. Run pip install -r requirements.txt first.")
    headers = {"User-Agent": "Mozilla/5.0 AI Content Summarizer"}
    response = requests.get(url, headers=headers, timeout=18)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "aside", "form"]):
        tag.decompose()
    title = soup.title.get_text(" ", strip=True) if soup.title else url
    content = " ".join(p.get_text(" ", strip=True) for p in soup.find_all(["h1", "h2", "h3", "p", "li"]))
    return title, clean_text(content)


def provider_client(provider: str, api_key: str):
    if provider == "Groq" and Groq and api_key:
        return Groq(api_key=api_key)
    if provider == "OpenAI" and OpenAI and api_key:
        return OpenAI(api_key=api_key)
    return None


def call_ai(
    text: str,
    provider: str,
    api_key: str,
    model: str,
    language: str,
    mode: str,
    summary_style: str,
    extra: str = "",
) -> str:
    client = provider_client(provider, api_key)
    if not client:
        return local_summary_pack(text, mode, summary_style)

    style_instruction = {
        "Executive Summary": "Create a concise executive summary with decisions, risks, and implications.",
        "Detailed Summary": "Create a structured detailed summary with headings.",
        "Bullet Insights": "Create bullet-point insights, each specific and non-redundant.",
        "Key Takeaways": "Create a short list of key takeaways.",
        "Action Items": "Extract action items with owners only when inferable.",
        "Topic Breakdown": "Break the content into topics with short explanations.",
        "One Sentence": "Summarize the entire content in exactly one sentence.",
        "Social Media": "Create a polished LinkedIn post, an X/Twitter post, and a short caption.",
        "Study Notes": "Convert the content into study notes with definitions, examples, and review questions.",
        "Presentation Points": "Create slide-ready presentation points with suggested slide titles.",
    }.get(summary_style, "Summarize the content clearly.")

    mode_instruction = (
        "Explain in beginner-friendly language with simple analogies and define jargon."
        if mode == "Beginner"
        else "Use expert-level framing, include nuance, assumptions, and strategic implications."
    )

    prompt = f"""
You are an AI research and productivity assistant.
Output language: {language}.
Explanation mode: {mode}.
Task: {style_instruction}
{mode_instruction}

Also include:
- Bullet-point insights
- Key takeaways
- Action items
- Topic-wise breakdown

{extra}

Content:
{text[:50000]}
"""
    try:
        if provider == "Groq":
            completion = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.25,
                max_tokens=2200,
            )
        else:
            completion = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.25,
                max_tokens=2200,
            )
        return completion.choices[0].message.content.strip()
    except Exception as exc:
        st.warning(f"AI provider failed, using local summary instead: {exc}")
        return local_summary_pack(text, mode, summary_style)


def local_summary_pack(text: str, mode: str, summary_style: str) -> str:
    short = extractive_summary(text, 4)
    one = extractive_summary(text, 1)
    kws = [kw for kw, _ in extract_keywords(text, 8)]
    topics = topic_breakdown(text, 4)
    action_sents = [s for s in sentences(text) if re.search(r"\b(should|must|need to|next|recommend|action|plan|implement|create|review)\b", s, re.I)]
    lines = [
        f"### {summary_style}",
        one if summary_style == "One Sentence" else short,
        "",
        "### Bullet-Point Insights",
        *[f"- {s}" for s in sentences(short)[:6]],
        "",
        "### Key Takeaways",
        *[f"- {kw.title()}" for kw in kws[:6]],
        "",
        "### Action Items",
        *([f"- {s}" for s in action_sents[:5]] or ["- No explicit action items were detected."]),
        "",
        f"### {mode} Explanation",
        "This version uses simpler language and highlights the core idea first." if mode == "Beginner" else "This version foregrounds structure, evidence, tradeoffs, and implications.",
        "",
        "### Topic-Wise Breakdown",
    ]
    for topic in topics:
        lines.append(f"- **{topic['topic']}**: {topic['summary']}")
    return "\n".join(lines)


def semantic_search(text: str, query: str, top_k: int = 5) -> pd.DataFrame:
    sents = sentences(text)
    if not query or not sents:
        return pd.DataFrame(columns=["score", "passage"])
    if TfidfVectorizer and cosine_similarity:
        vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        matrix = vectorizer.fit_transform(sents + [query])
        scores = cosine_similarity(matrix[-1], matrix[:-1]).flatten()
        rows = sorted(zip(scores, sents), reverse=True)[:top_k]
    else:
        q = set(words(query))
        rows = [(len(q.intersection(words(s))) / max(1, len(q)), s) for s in sents]
        rows = sorted(rows, reverse=True)[:top_k]
    return pd.DataFrame([{"score": round(float(score), 3), "passage": passage} for score, passage in rows])


def make_markdown_report(record: Dict, summary: str, metrics: Dict, keywords: List[Tuple[str, int]]) -> str:
    kw_text = ", ".join(k for k, _ in keywords[:15])
    return f"""# AI Content Summary Report

**Title:** {record.get('title', 'Untitled')}
**Generated:** {now_label()}
**Source Type:** {record.get('source_type', 'Content')}
**Language:** {metrics.get('language', 'Unknown')}

## Summary

{summary}

## Analytics

- Words: {metrics.get('words', 0)}
- Sentences: {metrics.get('sentences', 0)}
- Reading time: {metrics.get('reading_time', 0)} min
- Readability: {metrics.get('readability', 0)}
- Sentiment: {metrics.get('sentiment', 'Neutral')}

## Keywords

{kw_text}
"""


def simple_pdf_bytes(title: str, markdown_text: str) -> bytes:
    safe_lines = []
    for raw in markdown_text.replace("**", "").splitlines():
        line = re.sub(r"[^\x20-\x7E]", "", raw)
        if line.startswith("# "):
            safe_lines.extend(["", line[2:].upper(), ""])
        elif line.startswith("## "):
            safe_lines.extend(["", line[3:], ""])
        else:
            safe_lines.extend(textwrap.wrap(line, width=92) or [""])

    content = ["BT", "/F1 11 Tf", "50 780 Td", "14 TL"]
    for line in safe_lines[:260]:
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        content.append(f"({escaped}) Tj")
        content.append("T*")
    content.append("ET")
    stream = "\n".join(content).encode("latin-1", errors="ignore")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj",
        b"5 0 obj << /Length " + str(len(stream)).encode() + b" >> stream\n" + stream + b"\nendstream endobj",
    ]
    pdf = [b"%PDF-1.4\n"]
    offsets = []
    for obj in objects:
        offsets.append(sum(len(x) for x in pdf))
        pdf.append(obj + b"\n")
    xref = sum(len(x) for x in pdf)
    pdf.append(f"xref\n0 {len(objects)+1}\n0000000000 65535 f \n".encode())
    for off in offsets:
        pdf.append(f"{off:010d} 00000 n \n".encode())
    pdf.append(f"trailer << /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode())
    return b"".join(pdf)


def save_record(title: str, source_type: str, text: str, meta: Optional[Dict] = None) -> Dict:
    digest = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()[:16]
    record = {
        "id": digest,
        "title": title or "Untitled content",
        "source_type": source_type,
        "created_at": now_label(),
        "chars": len(text),
        "words": len(words(text)),
        "text": text,
        "meta": meta or {},
    }
    table = db().table("documents")
    existing = table.get(Query().id == digest)
    if existing:
        table.update(record, Query().id == digest)
    else:
        table.insert(record)
    return record


def load_history() -> List[Dict]:
    return sorted(db().table("documents").all(), key=lambda r: r.get("created_at", ""), reverse=True)


def tiny_html_icon(label: str, color: str) -> str:
    return f"<span class='badge'><span style='width:.55rem;height:.55rem;border-radius:50%;background:{color};display:inline-block'></span>{label}</span>"


def init_state():
    defaults = {
        "active_record": None,
        "summary": "",
        "chat_messages": [],
        "nav": "Summarizer",
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


init_state()


def safe_secret(name: str) -> str:
    try:
        return st.secrets.get(name, "")
    except Exception:
        return ""


def configure_langsmith(enabled: bool, api_key: str, project: str, endpoint: str = "") -> None:
    st.session_state["langsmith_enabled"] = bool(enabled and api_key and traceable)
    st.session_state["langsmith_project"] = project or "ai-content-summarizer"
    if enabled and api_key:
        os.environ["LANGSMITH_TRACING"] = "true"
        os.environ["LANGSMITH_API_KEY"] = api_key
        os.environ["LANGSMITH_PROJECT"] = project or "ai-content-summarizer"
        if endpoint:
            os.environ["LANGSMITH_ENDPOINT"] = endpoint
    else:
        os.environ["LANGSMITH_TRACING"] = "false"


def langsmith_trace(name: str, run_type: str, inputs: Dict, operation):
    if not st.session_state.get("langsmith_enabled") or not traceable:
        return operation()

    @traceable(name=name, run_type=run_type)
    def traced_operation(payload: Dict):
        return operation()

    return traced_operation(inputs)


with st.sidebar:
    st.markdown("## AI Summarizer")
    st.caption("Research, learning, and productivity workspace")
    nav = st.radio(
        "Navigation",
        ["Summarizer", "Analytics", "Semantic Search", "AI Chat", "History"],
        key="nav",
    )
    st.divider()
    provider = st.selectbox("AI Provider", ["Groq", "OpenAI", "Local fallback"])
    default_groq = safe_secret("GROQ_API_KEY") or os.getenv("GROQ_API_KEY", "")
    default_openai = safe_secret("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY", "")
    if provider == "Groq":
        api_key = st.text_input("Groq API Key", value=default_groq, type="password")
        model = st.text_input("Groq Model", value="llama-3.1-8b-instant")
    elif provider == "OpenAI":
        api_key = st.text_input("OpenAI API Key", value=default_openai, type="password")
        model = st.text_input("OpenAI Model", value="gpt-4o-mini")
    else:
        api_key = ""
        model = "local"
        st.info("Local mode uses extractive summaries and analytics.")
    st.divider()
    st.markdown("### LangSmith")
    default_langsmith = safe_secret("LANGSMITH_API_KEY") or os.getenv("LANGSMITH_API_KEY", "")
    langsmith_enabled = st.toggle(
        "Enable tracing",
        value=bool(default_langsmith),
        help="Send summary and chat runs to LangSmith for observability. Inputs and outputs may be visible in LangSmith.",
    )
    langsmith_key = st.text_input("LangSmith API Key", value=default_langsmith, type="password")
    langsmith_project = st.text_input(
        "LangSmith Project",
        value=safe_secret("LANGSMITH_PROJECT") or os.getenv("LANGSMITH_PROJECT", "ai-content-summarizer"),
    )
    langsmith_endpoint = st.text_input(
        "LangSmith Endpoint",
        value=safe_secret("LANGSMITH_ENDPOINT") or os.getenv("LANGSMITH_ENDPOINT", ""),
        placeholder="Optional, e.g. https://eu.api.smith.langchain.com",
    )
    configure_langsmith(langsmith_enabled, langsmith_key, langsmith_project, langsmith_endpoint)
    if langsmith_enabled and not traceable:
        st.warning("Install langsmith from requirements.txt to enable tracing.")
    elif st.session_state.get("langsmith_enabled"):
        st.success("LangSmith tracing is enabled.")
    else:
        st.caption("Tracing is off. The app still works normally.")
    st.divider()
    mode = st.radio("Explanation Mode", ["Beginner", "Expert"], horizontal=True)
    language = st.selectbox(
        "Output Language",
        ["English", "Hindi", "Telugu", "Spanish", "French", "German", "Arabic", "Chinese", "Japanese"],
    )
    st.caption("Supported sources: PDF, DOCX, TXT, pasted text, and webpages.")


st.markdown(
    """
<div class="hero">
  <h1>AI Powered Content Summarizer</h1>
  <p>Upload research papers, reports, notes, articles, webpages, or long-form text and turn them into sharp summaries, study notes, action plans, topic maps, and searchable knowledge.</p>
</div>
""",
    unsafe_allow_html=True,
)
st.write("")


def active_text() -> str:
    record = st.session_state.get("active_record")
    return record.get("text", "") if record else ""


def active_metrics(text: str) -> Dict:
    read = readability(text)
    sent = sentiment(text)
    return {
        "words": len(words(text)),
        "sentences": len(sentences(text)),
        "chars": len(text),
        "language": detect_language(text),
        "reading_time": read.get("Reading Time", 0),
        "readability": read.get("Flesch Reading Ease", 0),
        "grade": read.get("Grade Level", 0),
        "sentiment": sent.get("label", "Neutral"),
        "sentiment_score": sent.get("score", 0),
    }


if nav == "Summarizer":
    left, right = st.columns([1.05, 1.4], gap="large")
    with left:
        st.subheader("Add Content")
        source_mode = st.radio(
            "Content Source",
            ["Upload File", "Paste Text", "Webpage URL"],
            horizontal=True,
            help="Choose one input type so users do not need to fill every field.",
        )
        uploaded = None
        pasted = ""
        url = ""
        if source_mode == "Upload File":
            uploaded = st.file_uploader(
                "Drag and drop a PDF, DOCX, TXT, or Markdown file",
                type=["pdf", "docx", "txt", "md"],
                accept_multiple_files=False,
            )
            st.markdown('<div class="drop-hint">Supported file types: PDF research papers, DOCX reports, TXT notes, and Markdown drafts.</div>', unsafe_allow_html=True)
        elif source_mode == "Paste Text":
            pasted = st.text_area(
                "Paste article, notes, transcript, or long text",
                height=220,
                placeholder="Paste content here...",
            )
            st.markdown('<div class="drop-hint">Use this for articles, class notes, meeting transcripts, copied paper sections, or any long text.</div>', unsafe_allow_html=True)
        else:
            url = st.text_input("Webpage URL", placeholder="https://example.com/article")
            st.markdown('<div class="drop-hint">Use this only when you want the app to fetch and summarize a public webpage.</div>', unsafe_allow_html=True)

        ingest_col1, ingest_col2 = st.columns(2)
        with ingest_col1:
            ingest = st.button("Load Content", use_container_width=True)
        with ingest_col2:
            load_summary = st.button("Load & Summarize", use_container_width=True)
        clear = st.button("Clear Workspace", use_container_width=True)

        if clear:
            st.session_state.active_record = None
            st.session_state.summary = ""
            st.session_state.chat_messages = []
            st.rerun()

        if ingest or load_summary:
            text, title, source_type = "", "Untitled content", "Text"
            try:
                if source_mode == "Upload File" and uploaded:
                    data = uploaded.read()
                    suffix = uploaded.name.lower().split(".")[-1]
                    title = uploaded.name
                    source_type = suffix.upper()
                    if suffix == "pdf":
                        text = extract_pdf(data)
                    elif suffix == "docx":
                        text = extract_docx(data)
                    else:
                        text = extract_txt(data)
                    (UPLOAD_DIR / uploaded.name).write_bytes(data)
                elif source_mode == "Webpage URL" and url.strip():
                    title, text = fetch_webpage(url.strip())
                    source_type = "Webpage"
                elif source_mode == "Paste Text":
                    text = pasted
                    title = "Pasted content"
                    source_type = "Text"
                else:
                    st.error("Please add content for the selected source.")
                    text = ""
                text = clean_text(text)
                if not text:
                    st.error("No readable text was found. Try another file or paste content directly.")
                else:
                    st.session_state.active_record = save_record(title, source_type, text, {"url": url})
                    st.session_state.summary = ""
                    st.session_state.chat_messages = []
                    st.success(f"Loaded {len(words(text)):,} words from {title}.")
                    if load_summary:
                        with st.spinner("Loading and generating summary..."):
                            st.session_state.summary = langsmith_trace(
                                "load_and_summarize",
                                "chain",
                                {
                                    "title": title,
                                    "source_type": source_type,
                                    "provider": provider,
                                    "model": model,
                                    "language": language,
                                    "mode": mode,
                                    "style": "Executive Summary",
                                    "word_count": len(words(text)),
                                    "char_count": len(text),
                                },
                                lambda: call_ai(
                                    text,
                                    provider,
                                    api_key,
                                    model,
                                    language,
                                    mode,
                                    "Executive Summary",
                                ),
                            )
                            db().table("summaries").insert({
                                "document_id": st.session_state.active_record["id"],
                                "created_at": now_label(),
                                "style": "Executive Summary",
                                "mode": mode,
                                "language": language,
                                "summary": st.session_state.summary,
                            })
            except Exception as exc:
                st.error(f"Could not load content: {exc}")

    with right:
        record = st.session_state.active_record
        text = active_text()
        if record and text:
            metrics = active_metrics(text)
            c1, c2 = st.columns(2)
            c3, c4 = st.columns(2)
            cards = [
                ("Words", f"{metrics['words']:,}", "Content volume"),
                ("Read Time", f"{metrics['reading_time']} min", "At 220 wpm"),
                ("Language", metrics["language"], "Detected"),
                ("Sentiment", metrics["sentiment"], f"Score {metrics['sentiment_score']}"),
            ]
            for col, (label, value, note) in zip((c1, c2, c3, c4), cards):
                col.markdown(f"<div class='metric-card'><div class='metric-label'>{label}</div><div class='metric-value'>{value}</div><div class='metric-note'>{note}</div></div>", unsafe_allow_html=True)

            st.write("")
            tabs = st.tabs(["Summary Studio", "Standout Tools", "Report"])
            with tabs[0]:
                style = st.selectbox(
                    "Summary Type",
                    ["Executive Summary", "Detailed Summary", "Bullet Insights", "Key Takeaways", "Action Items", "Topic Breakdown"],
                )
                if st.button("Generate Summary", use_container_width=True):
                    with st.spinner("Thinking through the document..."):
                        st.session_state.summary = langsmith_trace(
                            "generate_summary",
                            "chain",
                            {
                                "title": record.get("title", "Untitled content"),
                                "source_type": record.get("source_type", "Content"),
                                "provider": provider,
                                "model": model,
                                "language": language,
                                "mode": mode,
                                "style": style,
                                "word_count": metrics["words"],
                                "char_count": metrics["chars"],
                            },
                            lambda: call_ai(text, provider, api_key, model, language, mode, style),
                        )
                        db().table("summaries").insert({
                            "document_id": record["id"],
                            "created_at": now_label(),
                            "style": style,
                            "mode": mode,
                            "language": language,
                            "summary": st.session_state.summary,
                        })
                if st.session_state.summary:
                    st.markdown(f"<div class='summary-box'>{st.session_state.summary}</div>", unsafe_allow_html=True)
                else:
                    st.info("Choose a summary type and generate your first AI output.")

            with tabs[1]:
                tool_cols = st.columns(2)
                tool_map = [
                    ("Summarize in one sentence", "One Sentence"),
                    ("Generate social media summary", "Social Media"),
                    ("Convert into study notes", "Study Notes"),
                    ("Explain for beginners", "Study Notes"),
                    ("Create presentation points automatically", "Presentation Points"),
                ]
                for idx, (label, style_name) in enumerate(tool_map):
                    if tool_cols[idx % 2].button(label, use_container_width=True):
                        with st.spinner("Creating polished output..."):
                            extra = "Make the explanation extra accessible for a beginner." if "beginners" in label else ""
                            trace_mode = "Beginner" if "beginners" in label else mode
                            st.session_state.summary = langsmith_trace(
                                "standout_summary_tool",
                                "chain",
                                {
                                    "tool": label,
                                    "title": record.get("title", "Untitled content"),
                                    "provider": provider,
                                    "model": model,
                                    "language": language,
                                    "mode": trace_mode,
                                    "style": style_name,
                                    "word_count": metrics["words"],
                                    "char_count": metrics["chars"],
                                },
                                lambda: call_ai(text, provider, api_key, model, language, trace_mode, style_name, extra),
                            )
                if st.session_state.summary:
                    st.markdown(f"<div class='summary-box'>{st.session_state.summary}</div>", unsafe_allow_html=True)

            with tabs[2]:
                keywords = extract_keywords(text)
                report_md = make_markdown_report(record, st.session_state.summary or extractive_summary(text, 5), metrics, keywords)
                st.download_button("Download Markdown Report", report_md, file_name="summary_report.md", mime="text/markdown", use_container_width=True)
                st.download_button("Download PDF Report", simple_pdf_bytes(record["title"], report_md), file_name="summary_report.pdf", mime="application/pdf", use_container_width=True)
                st.code(report_md[:1800] + ("..." if len(report_md) > 1800 else ""), language="markdown")
        else:
            st.markdown('<div class="glass">Load a file, webpage, or pasted text to unlock summaries, analytics, reports, search, and chat.</div>', unsafe_allow_html=True)


elif nav == "Analytics":
    text = active_text()
    if not text:
        st.info("Load content in the Summarizer tab first.")
    else:
        metrics = active_metrics(text)
        keywords = extract_keywords(text, 20)
        c1, c2, c3, c4 = st.columns(4)
        for col, (label, value, note) in zip(
            (c1, c2, c3, c4),
            [
                ("Characters", f"{metrics['chars']:,}", "Raw length"),
                ("Sentences", f"{metrics['sentences']:,}", "Detected"),
                ("Readability", metrics["readability"], "Flesch score"),
                ("Grade", metrics["grade"], "Approximate"),
            ],
        ):
            col.markdown(f"<div class='metric-card'><div class='metric-label'>{label}</div><div class='metric-value'>{value}</div><div class='metric-note'>{note}</div></div>", unsafe_allow_html=True)

        left, right = st.columns(2, gap="large")
        with left:
            kw_df = pd.DataFrame(keywords, columns=["Keyword", "Frequency"])
            fig = px.bar(kw_df, x="Frequency", y="Keyword", orientation="h", color="Frequency", color_continuous_scale=["#36d1dc", "#a7f070"])
            fig.update_layout(template="plotly_dark", height=460, margin=dict(l=10, r=10, t=30, b=10), yaxis=dict(autorange="reversed"))
            st.plotly_chart(fig, use_container_width=True)
        with right:
            sent = sentiment(text)
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=(sent["score"] + 1) * 50,
                title={"text": f"Sentiment: {sent['label']}"},
                gauge={"axis": {"range": [0, 100]}, "bar": {"color": "#36d1dc"}, "steps": [
                    {"range": [0, 40], "color": "rgba(255,113,154,.35)"},
                    {"range": [40, 60], "color": "rgba(244,201,93,.35)"},
                    {"range": [60, 100], "color": "rgba(167,240,112,.35)"},
                ]},
            ))
            fig.update_layout(template="plotly_dark", height=320, margin=dict(l=20, r=20, t=50, b=20))
            st.plotly_chart(fig, use_container_width=True)
            lengths = [len(words(s)) for s in sentences(text)]
            hist = px.histogram(pd.DataFrame({"Sentence Length": lengths}), x="Sentence Length", nbins=24, color_discrete_sequence=["#f4c95d"])
            hist.update_layout(template="plotly_dark", height=260, margin=dict(l=10, r=10, t=30, b=10))
            st.plotly_chart(hist, use_container_width=True)

        st.subheader("Topic Map")
        topic_df = pd.DataFrame(topic_breakdown(text, 7))
        st.dataframe(topic_df, use_container_width=True, hide_index=True)


elif nav == "Semantic Search":
    text = active_text()
    if not text:
        st.info("Load content in the Summarizer tab first.")
    else:
        query = st.text_input("Search inside your content", placeholder="Find claims, methods, risks, examples, next steps...")
        results = semantic_search(text, query)
        if query:
            st.dataframe(results, use_container_width=True, hide_index=True)
        else:
            st.markdown('<div class="glass">Ask a semantic question and the app will rank the most relevant passages.</div>', unsafe_allow_html=True)


elif nav == "AI Chat":
    text = active_text()
    if not text:
        st.info("Load content in the Summarizer tab first.")
    else:
        for msg in st.session_state.chat_messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
        question = st.chat_input("Ask anything about the uploaded content")
        if question:
            st.session_state.chat_messages.append({"role": "user", "content": question})
            with st.chat_message("user"):
                st.markdown(question)
            with st.chat_message("assistant"):
                with st.spinner("Reading the document context..."):
                    context_df = semantic_search(text, question, 6)
                    context = "\n".join(context_df["passage"].tolist())
                    answer = langsmith_trace(
                        "content_chat",
                        "chain",
                        {
                            "question": question,
                            "provider": provider,
                            "model": model,
                            "language": language,
                            "mode": mode,
                            "retrieved_passages": len(context_df),
                            "context_chars": len(context),
                        },
                        lambda: call_ai(
                            context or text[:8000],
                            provider,
                            api_key,
                            model,
                            language,
                            mode,
                            "Detailed Summary",
                            extra=f"Answer this user question using only the supplied content: {question}",
                        ),
                    )
                    st.markdown(answer)
            st.session_state.chat_messages.append({"role": "assistant", "content": answer})


elif nav == "History":
    history = load_history()
    if not history:
        st.info("No uploads yet.")
    else:
        st.subheader("Upload and Analysis History")
        for item in history:
            with st.expander(f"{item['title']} | {item['source_type']} | {item['created_at']}"):
                st.write(f"Words: {item.get('words', 0):,} | Characters: {item.get('chars', 0):,}")
                if st.button("Open This Document", key=f"open-{item['id']}"):
                    st.session_state.active_record = item
                    st.session_state.summary = ""
                    st.session_state.nav = "Summarizer"
                    st.rerun()
                st.text_area("Preview", item.get("text", "")[:2500], height=160, key=f"preview-{item['id']}")
